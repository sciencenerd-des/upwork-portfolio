"""
Report Builder Module for the Automated Report Generator.

This module handles:
- PDF report generation using ReportLab
- Word document generation using python-docx
- Professional formatting with headers, footers, and page numbers
- Embedding charts and tables in reports
"""

import io
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
from datetime import datetime

# ReportLab imports for PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4, LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle,
    PageBreak, KeepTogether, ListFlowable, ListItem
)
from reportlab.platypus.flowables import HRFlowable

# python-docx imports for Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import pandas as pd

from .utils import (
    get_styles_config,
    hex_to_rgb,
    format_currency,
    format_number,
    format_percentage,
    generate_report_filename,
    ensure_directory,
)


class ReportBuilderError(Exception):
    """Exception raised for report building errors."""
    pass


class ReportBuilder:
    """
    Builds professional PDF and Word reports.

    Handles document structure, styling, and content embedding
    based on configuration from styles.yaml.
    """

    def __init__(self):
        """Initialize the ReportBuilder with styling configuration."""
        self._styles = get_styles_config()
        self._pdf_config = self._styles.get("pdf", {})
        self._word_config = self._styles.get("word", {})
        self._colors = self._styles.get("colors", {})
        self._table_config = self._styles.get("tables", {})

    def _get_page_size(self, config_page_size: str) -> tuple:
        """Get page size tuple from string."""
        sizes = {
            "LETTER": LETTER,
            "A4": A4,
            "LEGAL": (8.5 * inch, 14 * inch),
        }
        return sizes.get(config_page_size.upper(), LETTER)

    def _hex_to_reportlab_color(self, hex_color: str) -> colors.Color:
        """Convert hex color to ReportLab color."""
        rgb = hex_to_rgb(hex_color)
        return colors.Color(rgb[0]/255, rgb[1]/255, rgb[2]/255)

    def _hex_to_docx_color(self, hex_color: str) -> RGBColor:
        """Convert hex color to python-docx RGBColor."""
        rgb = hex_to_rgb(hex_color)
        return RGBColor(rgb[0], rgb[1], rgb[2])

    # ==================== PDF Generation ====================

    def build_pdf(
        self,
        output_path: Union[str, Path],
        title: str,
        sections: List[Dict[str, Any]],
        metadata: Optional[Dict[str, str]] = None,
    ) -> str:
        """
        Build a PDF report.

        Args:
            output_path: Path for the output PDF file
            title: Report title
            sections: List of section dictionaries with content
            metadata: Optional metadata (author, date, etc.)

        Returns:
            Path to the generated PDF file
        """
        output_path = Path(output_path)
        ensure_directory(output_path.parent)

        # Get page settings
        page_size = self._get_page_size(self._pdf_config.get("page_size", "LETTER"))
        margins = self._pdf_config.get("margins", {})

        # Create document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=page_size,
            topMargin=margins.get("top", 72),
            bottomMargin=margins.get("bottom", 72),
            leftMargin=margins.get("left", 72),
            rightMargin=margins.get("right", 72),
        )

        # Build story (content)
        story = []
        styles = self._create_pdf_styles()

        # Title
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, self._pdf_config.get("spacing", {}).get("after_title", 24)))

        # Metadata line
        if metadata:
            meta_text = f"Generated: {metadata.get('date', datetime.now().strftime('%B %d, %Y'))}"
            if metadata.get('period'):
                meta_text += f" | Period: {metadata['period']}"
            story.append(Paragraph(meta_text, styles['Caption']))
            story.append(Spacer(1, 12))

        # Horizontal line
        story.append(HRFlowable(
            width="100%",
            thickness=1,
            color=self._hex_to_reportlab_color(self._colors.get("neutral", "#6B7280")),
            spaceBefore=6,
            spaceAfter=12
        ))

        # Process sections
        for section in sections:
            section_content = self._build_pdf_section(section, styles)
            story.extend(section_content)

        # Build PDF with page numbers
        doc.build(
            story,
            onFirstPage=self._pdf_header_footer,
            onLaterPages=self._pdf_header_footer
        )

        return str(output_path)

    def _create_pdf_styles(self) -> Dict[str, ParagraphStyle]:
        """Create PDF paragraph styles from configuration."""
        base_styles = getSampleStyleSheet()
        custom_styles = {}

        font_config = self._pdf_config.get("fonts", {})

        # Title style
        title_cfg = font_config.get("title", {})
        custom_styles['Title'] = ParagraphStyle(
            'CustomTitle',
            parent=base_styles['Title'],
            fontName=title_cfg.get("name", "Helvetica-Bold"),
            fontSize=title_cfg.get("size", 24),
            textColor=self._hex_to_reportlab_color(title_cfg.get("color", "#111827")),
            alignment=TA_CENTER,
            spaceAfter=12,
        )

        # Heading1 style
        h1_cfg = font_config.get("heading1", {})
        custom_styles['Heading1'] = ParagraphStyle(
            'CustomH1',
            parent=base_styles['Heading1'],
            fontName=h1_cfg.get("name", "Helvetica-Bold"),
            fontSize=h1_cfg.get("size", 18),
            textColor=self._hex_to_reportlab_color(h1_cfg.get("color", "#111827")),
            spaceBefore=18,
            spaceAfter=12,
        )

        # Heading2 style
        h2_cfg = font_config.get("heading2", {})
        custom_styles['Heading2'] = ParagraphStyle(
            'CustomH2',
            parent=base_styles['Heading2'],
            fontName=h2_cfg.get("name", "Helvetica-Bold"),
            fontSize=h2_cfg.get("size", 14),
            textColor=self._hex_to_reportlab_color(h2_cfg.get("color", "#374151")),
            spaceBefore=12,
            spaceAfter=8,
        )

        # Body style
        body_cfg = font_config.get("body", {})
        custom_styles['Body'] = ParagraphStyle(
            'CustomBody',
            parent=base_styles['Normal'],
            fontName=body_cfg.get("name", "Helvetica"),
            fontSize=body_cfg.get("size", 11),
            textColor=self._hex_to_reportlab_color(body_cfg.get("color", "#4B5563")),
            alignment=TA_JUSTIFY,
            spaceBefore=4,
            spaceAfter=8,
        )

        # Caption style
        caption_cfg = font_config.get("caption", {})
        custom_styles['Caption'] = ParagraphStyle(
            'CustomCaption',
            parent=base_styles['Normal'],
            fontName=caption_cfg.get("name", "Helvetica-Oblique"),
            fontSize=caption_cfg.get("size", 9),
            textColor=self._hex_to_reportlab_color(caption_cfg.get("color", "#6B7280")),
            alignment=TA_CENTER,
        )

        # Bullet style
        custom_styles['Bullet'] = ParagraphStyle(
            'CustomBullet',
            parent=custom_styles['Body'],
            leftIndent=20,
            bulletIndent=10,
        )

        return custom_styles

    def _build_pdf_section(
        self,
        section: Dict[str, Any],
        styles: Dict[str, ParagraphStyle]
    ) -> List:
        """Build a section for the PDF document."""
        content = []
        section_type = section.get("type", "text")

        # Section title
        if section.get("title"):
            content.append(Paragraph(section["title"], styles['Heading1']))

        if section_type == "summary":
            content.extend(self._build_pdf_summary(section, styles))
        elif section_type == "chart":
            content.extend(self._build_pdf_chart(section, styles))
        elif section_type == "table":
            content.extend(self._build_pdf_table(section, styles))
        elif section_type == "insights":
            content.extend(self._build_pdf_insights(section, styles))
        elif section_type == "text":
            content.extend(self._build_pdf_text(section, styles))

        # Add spacing after section
        spacing = self._pdf_config.get("spacing", {})
        content.append(Spacer(1, spacing.get("after_paragraph", 8)))

        return content

    def _build_pdf_summary(
        self,
        section: Dict[str, Any],
        styles: Dict[str, ParagraphStyle]
    ) -> List:
        """Build summary metrics section for PDF."""
        content = []
        metrics = section.get("metrics", [])

        if not metrics:
            return content

        # Create a table for metrics
        table_data = []
        row = []

        for i, metric in enumerate(metrics):
            cell_content = [
                Paragraph(f"<b>{metric.get('value', '')}</b>", styles['Body']),
                Paragraph(metric.get('label', ''), styles['Caption']),
            ]
            row.append(cell_content)

            # 4 metrics per row
            if (i + 1) % 4 == 0 or i == len(metrics) - 1:
                table_data.append(row)
                row = []

        if table_data:
            # Pad last row if needed
            while len(table_data[-1]) < 4:
                table_data[-1].append([
                    Paragraph('', styles['Body']),
                    Paragraph('', styles['Caption']),
                ])

            table = Table(table_data, colWidths=[1.5*inch] * 4)
            table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BACKGROUND', (0, 0), (-1, -1), self._hex_to_reportlab_color(
                    self._styles.get("summary_cards", {}).get("background_color", "#F3F4F6")
                )),
                ('TOPPADDING', (0, 0), (-1, -1), 12),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ]))
            content.append(table)

        return content

    def _build_pdf_chart(
        self,
        section: Dict[str, Any],
        styles: Dict[str, ParagraphStyle]
    ) -> List:
        """Build chart section for PDF."""
        content = []
        image_data = section.get("image_bytes")
        image_path = section.get("image_path")

        if image_data:
            img = Image(io.BytesIO(image_data), width=6*inch, height=3.6*inch)
            content.append(img)
        elif image_path and Path(image_path).exists():
            img = Image(str(image_path), width=6*inch, height=3.6*inch)
            content.append(img)

        if section.get("caption"):
            content.append(Spacer(1, 6))
            content.append(Paragraph(section["caption"], styles['Caption']))

        spacing = self._pdf_config.get("spacing", {})
        content.append(Spacer(1, spacing.get("after_chart", 18)))

        return content

    def _build_pdf_table(
        self,
        section: Dict[str, Any],
        styles: Dict[str, ParagraphStyle]
    ) -> List:
        """Build data table section for PDF."""
        content = []
        df = section.get("dataframe")

        if df is None or len(df) == 0:
            return content

        # Limit rows
        max_rows = self._table_config.get("max_rows_display", 20)
        if len(df) > max_rows:
            df = df.head(max_rows)

        # Convert DataFrame to table data
        table_data = [list(df.columns)]  # Headers
        for _, row in df.iterrows():
            formatted_row = []
            for col, val in row.items():
                if pd.isna(val):
                    formatted_row.append("")
                elif isinstance(val, float):
                    formatted_row.append(format_number(val, 2))
                elif hasattr(val, 'strftime'):
                    # Format datetime/date objects as date only
                    formatted_row.append(val.strftime('%Y-%m-%d'))
                else:
                    formatted_row.append(str(val))
            table_data.append(formatted_row)

        # Create table
        col_width = (6 * inch) / len(df.columns)
        table = Table(table_data, colWidths=[col_width] * len(df.columns))

        # Style table
        header_color = self._hex_to_reportlab_color(self._colors.get("primary", "#2563EB"))
        table.setStyle(TableStyle([
            # Header
            ('BACKGROUND', (0, 0), (-1, 0), header_color),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            # Body
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            # Grid
            ('GRID', (0, 0), (-1, -1), 0.5, self._hex_to_reportlab_color("#E5E7EB")),
            # Alternating rows
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white,
             self._hex_to_reportlab_color("#F9FAFB")]),
        ]))

        content.append(table)
        spacing = self._pdf_config.get("spacing", {})
        content.append(Spacer(1, spacing.get("after_table", 12)))

        return content

    def _build_pdf_insights(
        self,
        section: Dict[str, Any],
        styles: Dict[str, ParagraphStyle]
    ) -> List:
        """Build insights section for PDF."""
        content = []
        insights = section.get("insights", [])

        if not insights:
            return content

        for insight in insights:
            bullet_text = f"<bullet>&bull;</bullet> {insight}"
            content.append(Paragraph(bullet_text, styles['Bullet']))

        return content

    def _build_pdf_text(
        self,
        section: Dict[str, Any],
        styles: Dict[str, ParagraphStyle]
    ) -> List:
        """Build text section for PDF."""
        content = []
        text = section.get("text", "")

        if text:
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    content.append(Paragraph(para.strip(), styles['Body']))

        return content

    def _pdf_header_footer(self, canvas, doc):
        """Add header and footer to PDF pages."""
        canvas.saveState()
        page_width, page_height = doc.pagesize

        footer_config = self._pdf_config.get("footer", {})

        # Footer
        if footer_config.get("show_page_number", True):
            page_num = canvas.getPageNumber()
            footer_text = f"Page {page_num}"

            if footer_config.get("show_generated_by", True):
                footer_text = f"{footer_config.get('text', 'Generated by Automated Report Generator')} | {footer_text}"

            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(self._hex_to_reportlab_color(self._colors.get("neutral", "#6B7280")))
            canvas.drawCentredString(page_width / 2, 30, footer_text)

        canvas.restoreState()

    # ==================== Word Document Generation ====================

    def build_word(
        self,
        output_path: Union[str, Path],
        title: str,
        sections: List[Dict[str, Any]],
        metadata: Optional[Dict[str, str]] = None,
    ) -> str:
        """
        Build a Word document report.

        Args:
            output_path: Path for the output DOCX file
            title: Report title
            sections: List of section dictionaries with content
            metadata: Optional metadata (author, date, etc.)

        Returns:
            Path to the generated Word document
        """
        output_path = Path(output_path)
        ensure_directory(output_path.parent)

        doc = Document()

        # Set up styles
        self._setup_word_styles(doc)

        # Set margins
        margins = self._word_config.get("margins", {})
        for section in doc.sections:
            section.top_margin = Inches(margins.get("top", 1.0))
            section.bottom_margin = Inches(margins.get("bottom", 1.0))
            section.left_margin = Inches(margins.get("left", 1.0))
            section.right_margin = Inches(margins.get("right", 1.0))

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if metadata:
            meta_text = f"Generated: {metadata.get('date', datetime.now().strftime('%B %d, %Y'))}"
            if metadata.get('period'):
                meta_text += f" | Period: {metadata['period']}"
            meta_para = doc.add_paragraph(meta_text)
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in meta_para.runs:
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = self._hex_to_docx_color("#6B7280")

        # Add horizontal line
        doc.add_paragraph()
        self._add_word_horizontal_line(doc)

        # Process sections
        for section in sections:
            self._build_word_section(doc, section)

        # Save document
        doc.save(str(output_path))

        return str(output_path)

    def _setup_word_styles(self, doc: Document):
        """Set up custom styles for Word document."""
        styles = doc.styles
        font_config = self._word_config.get("fonts", {})

        # Modify heading styles
        try:
            h1_style = styles['Heading 1']
            h1_cfg = font_config.get("heading1", {})
            h1_style.font.name = h1_cfg.get("name", "Arial")
            h1_style.font.size = Pt(h1_cfg.get("size", 18))
            h1_style.font.bold = h1_cfg.get("bold", True)
            h1_style.font.color.rgb = self._hex_to_docx_color(h1_cfg.get("color", "111827"))
        except Exception:
            pass

    def _add_word_horizontal_line(self, doc: Document):
        """Add a horizontal line to Word document."""
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_after = Pt(12)

        # Create border element
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '6B7280')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _build_word_section(self, doc: Document, section: Dict[str, Any]):
        """Build a section for the Word document."""
        section_type = section.get("type", "text")

        # Section title
        if section.get("title"):
            doc.add_heading(section["title"], level=1)

        if section_type == "summary":
            self._build_word_summary(doc, section)
        elif section_type == "chart":
            self._build_word_chart(doc, section)
        elif section_type == "table":
            self._build_word_table(doc, section)
        elif section_type == "insights":
            self._build_word_insights(doc, section)
        elif section_type == "text":
            self._build_word_text(doc, section)

        # Add spacing
        doc.add_paragraph()

    def _build_word_summary(self, doc: Document, section: Dict[str, Any]):
        """Build summary metrics section for Word."""
        metrics = section.get("metrics", [])

        if not metrics:
            return

        # Create a table for metrics (2 columns)
        num_cols = min(4, len(metrics))
        num_rows = (len(metrics) + num_cols - 1) // num_cols

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table_config = self._word_config.get("table", {})

        for i, metric in enumerate(metrics):
            row_idx = i // num_cols
            col_idx = i % num_cols
            cell = table.cell(row_idx, col_idx)

            # Value
            p_value = cell.paragraphs[0]
            run_value = p_value.add_run(str(metric.get('value', '')))
            run_value.font.size = Pt(16)
            run_value.font.bold = True
            p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Label
            p_label = cell.add_paragraph()
            run_label = p_label.add_run(metric.get('label', ''))
            run_label.font.size = Pt(9)
            run_label.font.color.rgb = self._hex_to_docx_color("#6B7280")
            p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Cell shading
            cell_shading = OxmlElement('w:shd')
            cell_shading.set(qn('w:fill'), 'F3F4F6')
            cell._tc.get_or_add_tcPr().append(cell_shading)

    def _build_word_chart(self, doc: Document, section: Dict[str, Any]):
        """Build chart section for Word."""
        image_data = section.get("image_bytes")
        image_path = section.get("image_path")

        if image_data:
            # Add image from bytes
            image_stream = io.BytesIO(image_data)
            doc.add_picture(image_stream, width=Inches(6))
        elif image_path and Path(image_path).exists():
            doc.add_picture(str(image_path), width=Inches(6))

        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add caption
        if section.get("caption"):
            caption = doc.add_paragraph(section["caption"])
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                run.font.size = Pt(9)
                run.font.italic = True

    def _build_word_table(self, doc: Document, section: Dict[str, Any]):
        """Build data table section for Word."""
        df = section.get("dataframe")

        if df is None or len(df) == 0:
            return

        # Limit rows
        max_rows = self._table_config.get("max_rows_display", 20)
        if len(df) > max_rows:
            df = df.head(max_rows)

        # Create table
        table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
        table.style = 'Table Grid'

        table_config = self._word_config.get("table", {})

        # Header row
        header_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            header_cells[i].text = str(col_name)
            # Style header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = self._hex_to_docx_color(
                        table_config.get("header_text_color", "FFFFFF")
                    )

            # Header background
            cell_shading = OxmlElement('w:shd')
            cell_shading.set(qn('w:fill'), table_config.get("header_background", "2563EB"))
            header_cells[i]._tc.get_or_add_tcPr().append(cell_shading)

        # Data rows
        for row_idx, (_, row) in enumerate(df.iterrows()):
            cells = table.rows[row_idx + 1].cells
            for col_idx, val in enumerate(row):
                if pd.isna(val):
                    cells[col_idx].text = ""
                elif isinstance(val, float):
                    cells[col_idx].text = format_number(val, 2)
                elif hasattr(val, 'strftime'):
                    # Format datetime/date objects as date only
                    cells[col_idx].text = val.strftime('%Y-%m-%d')
                else:
                    cells[col_idx].text = str(val)

            # Alternating row colors
            if row_idx % 2 == 0:
                for cell in cells:
                    cell_shading = OxmlElement('w:shd')
                    cell_shading.set(qn('w:fill'), table_config.get("row_even_background", "F9FAFB"))
                    cell._tc.get_or_add_tcPr().append(cell_shading)

    def _build_word_insights(self, doc: Document, section: Dict[str, Any]):
        """Build insights section for Word."""
        insights = section.get("insights", [])

        if not insights:
            return

        for insight in insights:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(insight)

    def _build_word_text(self, doc: Document, section: Dict[str, Any]):
        """Build text section for Word."""
        text = section.get("text", "")

        if text:
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

    # ==================== Convenience Methods ====================

    def build_report(
        self,
        output_dir: Union[str, Path],
        title: str,
        sections: List[Dict[str, Any]],
        template_name: str,
        formats: List[str] = ["pdf", "docx"],
        metadata: Optional[Dict[str, str]] = None,
    ) -> Dict[str, str]:
        """
        Build reports in multiple formats.

        Args:
            output_dir: Directory for output files
            title: Report title
            sections: List of section dictionaries
            template_name: Name of the template (for filename)
            formats: List of formats to generate
            metadata: Optional metadata

        Returns:
            Dictionary mapping format to file path
        """
        output_dir = Path(output_dir)
        ensure_directory(output_dir)

        results = {}

        for fmt in formats:
            filename = generate_report_filename(template_name, fmt)
            output_path = output_dir / filename

            if fmt == "pdf":
                results["pdf"] = self.build_pdf(output_path, title, sections, metadata)
            elif fmt in ["docx", "word"]:
                results["docx"] = self.build_word(output_path, title, sections, metadata)

        return results
